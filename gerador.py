import xlrd, datetime
from docx import Document
from docx.shared import Inches
from docx.text.run import Font, Run
from docx.shared import RGBColor

while True:
    file_name_original = input('Digite o nome do arquivo original: ')
    if file_name_original != '':
        break
while True:
    city = input('Digite o nome da cidade: ')
    if city != '':
        break
while True:
    farm_name = input('Digite o nome da fazenda: ')
    if farm_name != '':
        break

employes  = xlrd.open_workbook(file_name_original+ '.xls')
sheet = employes.sheet_by_index(0)

for row in range(4, sheet.nrows):
    employe_cod = str( sheet.cell(row, 0).value )
    employe_name = sheet.cell(row, 1).value
    employe_data_adt = int( sheet.cell(row, 2).value )
    employe_ctps = str( sheet.cell( row, 3 ).value ).split("/")[0]
    employe_serie = str( sheet.cell( row, 3 ).value ).split("/")[1]
    employe_data_adt = str( datetime.date(1900, 1, 1) + datetime.timedelta(employe_data_adt - 2) )
    year = employe_data_adt.split("-")[0]
    month = employe_data_adt.split("-")[1]
    day = employe_data_adt.split("-")[2]
    document = Document()
    p = document.add_heading('DECLARAÇÃO', 0)
    paragraph_1 = f'''    Eu {employe_name} portador(a) da CTPS nº. {employe_ctps} Série {employe_serie} funcionário(a) da Fazenda {farm_name} no setor de colheita, DECLARO EXPRESSAMENTE ter pleno conhecimento de que o intervalo para as refeições é no horário compreendido das 11:00 até 12:00 hs, e que há na empresa determinação expressa no sentido de que referido intervalo deve ser inteiramente respeitado, ou seja, não se deve realizar a colheita no referido horário. '''
    paragraph_2 = ''' 				Declaro ainda, estar plenamente ciente de que nenhuma produção será anotada no referido intervalo. '''
    paragraph_3 = f''' {city}/SP, {day}/{month}/{year}. '''
    paragraph_4 = '''_________________________________ '''
    paragraph_5 = '''Assinatura do empregado   '''
    p1 = document.add_paragraph(paragraph_1)
    p2 = document.add_paragraph(paragraph_2)
    p3 = document.add_paragraph(paragraph_3)
    p4 = document.add_paragraph(paragraph_4)
    p5 = document.add_paragraph(paragraph_5)
    p.alignment = 1
    p4.alignment = 1
    p5.alignment = 1

    print( f"Cod: {employe_cod} | Nome: {employe_name} ")
    archive_name = r'funcionario_' + employe_cod + '.docx'
    document.save(archive_name)
    
